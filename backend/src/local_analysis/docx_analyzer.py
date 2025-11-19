"""
DOCX Document Analyzer Extension
Handles Microsoft Word (.docx) files
"""
from pathlib import Path
from typing import Optional, Tuple, TYPE_CHECKING
import logging

if TYPE_CHECKING:
    from docx import Document as DocxDocument

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None  # type: ignore

from document_analyzer import (
    DocumentAnalyzer, DocumentConfig, DocumentAnalysisResult,
    DocumentMetadata
)

logger = logging.getLogger(__name__)


class DocxAnalyzer(DocumentAnalyzer):
    """
    Extended DocumentAnalyzer with .docx support
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.markdown', '.rst', '.log', '.docx'}
    
    def __init__(self, config: Optional[DocumentConfig] = None):
        """Initialize with .docx support check"""
        super().__init__(config)
        
        if not DOCX_AVAILABLE:
            logger.warning(
                "python-docx not installed. .docx files will not be supported. "
                "Install with: pip install python-docx"
            )
    
    def _extract_docx_content(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """
        Extract content and metadata from .docx file
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Tuple of (content, metadata)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for .docx support")
        
        from docx import Document as DocxDocument  # Import here to satisfy type checker
        doc = DocxDocument(str(file_path))
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)
        
        # Calculate basic metadata
        word_count = len(content.split())
        character_count = len(content)
        line_count = content.count('\n') + 1
        paragraph_count = len(paragraphs)
        reading_time_minutes = word_count / 200.0
        
        # DOCX-specific metadata
        sections = len(doc.sections)
        
        # Estimate pages (rough calculation: ~250 words per page)
        estimated_pages = max(1, word_count // 250)
        
        metadata = DocumentMetadata(
            file_type='.docx',
            word_count=word_count,
            character_count=character_count,
            line_count=line_count,
            paragraph_count=paragraph_count,
            reading_time_minutes=reading_time_minutes,
            sections=sections,
            pages=estimated_pages
        )
        
        # Extract headings (from styles)
        headings = []
        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text)
        
        metadata.headings = headings
        metadata.heading_count = len(headings)
        
        return content, metadata
    
    def analyze_document(self, file_path: Path) -> DocumentAnalysisResult:
        """
        Analyze document (with .docx support)
        
        Args:
            file_path: Path to document file
            
        Returns:
            DocumentAnalysisResult
        """
        file_type = file_path.suffix.lower()
        
        # Handle .docx files specially
        if file_type == '.docx':
            if not DOCX_AVAILABLE:
                return DocumentAnalysisResult(
                    file_name=file_path.name,
                    file_type=file_type,
                    success=False,
                    error_message="python-docx library not installed"
                )
            
            return self._analyze_docx(file_path)
        
        # Use parent class for other file types
        return super().analyze_document(file_path)
    
    def _analyze_docx(self, file_path: Path) -> DocumentAnalysisResult:
        """
        Analyze .docx file
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            DocumentAnalysisResult
        """
        result = DocumentAnalysisResult(
            file_name=file_path.name,
            file_type='.docx',
            success=False
        )
        
        try:
            # Validate file size
            is_valid, file_size_mb, error_msg = self.validate_file_size(file_path)
            result.file_size_mb = file_size_mb
            
            if not is_valid:
                result.error_message = error_msg
                return result
            
            # Extract content and metadata
            content, metadata = self._extract_docx_content(file_path)
            result.text_content = content
            result.metadata = metadata
            
            if not content.strip():
                result.error_message = "Empty document"
                return result
            
            # Generate summary
            summary_result = self.summarizer.generate_summary(content, file_path.name)
            
            if summary_result.success:
                result.summary = summary_result.summary_text
                result.keywords = summary_result.keywords
                result.statistics = summary_result.statistics
            else:
                logger.warning(f"Summary generation failed for {file_path.name}")
                result.summary = ""
                result.keywords = []
                result.statistics = {}
            
            # Extract key topics
            result.key_topics = self._extract_key_topics(content, metadata)
            
            result.success = True
            logger.info(
                f"Successfully analyzed {file_path.name}: "
                f"{metadata.word_count} words, "
                f"{metadata.pages} pages (estimated)"
            )
            
        except Exception as e:
            error_msg = f"Error analyzing .docx file: {str(e)}"
            logger.error(f"{file_path.name}: {error_msg}")
            result.error_message = error_msg
        
        return result
    
    def to_json(self, result: DocumentAnalysisResult) -> dict:
        """
        Convert result to JSON (with .docx metadata)
        
        Args:
            result: DocumentAnalysisResult
            
        Returns:
            Dictionary with results
        """
        json_data = super().to_json(result)
        
        # Add .docx-specific metadata
        if result.success and result.file_type == '.docx' and result.metadata:
            json_data['metadata']['docx'] = {
                'pages': result.metadata.pages,
                'sections': result.metadata.sections,
                'headings': result.metadata.headings,
            }
        
        return json_data


def create_docx_analyzer(
    max_file_size_mb: float = 10.0,
    max_batch_size: int = 20
) -> DocxAnalyzer:
    """
    Factory function to create a DocxAnalyzer
    
    Args:
        max_file_size_mb: Maximum file size in MB
        max_batch_size: Maximum documents per batch
        
    Returns:
        Configured DocxAnalyzer instance
    """
    config = DocumentConfig(
        max_file_size_mb=max_file_size_mb,
        max_batch_size=max_batch_size
    )
    return DocxAnalyzer(config)
