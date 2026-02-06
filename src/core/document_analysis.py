"""Lightweight, non-LLM document analysis for project workspaces.
Supported formats: docx, pdf, txt, md. Outputs per-file summaries with hashes
for deduplication plus simple heuristics for roles, dates, metrics, and skills."""
from __future__ import annotations
import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional
try:
    from docx import Document  # type: ignore
except Exception:
    Document = None
try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    PdfReader = None
SUPPORTED_DOC_EXTS = {".docx", ".pdf", ".txt", ".md"}

def compute_sha256(path: Path) -> str:
    """
    Compute a SHA256 hash of a file for deduplication.
    Args: path (Path): File path to hash.
    Returns: str: Hex digest of the file content.
    """
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            digest.update(chunk)
    return digest.hexdigest()
@dataclass

class ParsedDoc:
    text: str
    headings: List[str] = field(default_factory=list)
    page_count: int | None = None

class DocumentAnalyzer:
    """
    Analyze text-bearing project documents without external services.
    Provides:
      - Content hashing for deduplication
      - Shallow text extraction per format
      - Heuristic extraction of metrics, dates, roles, and skills
    """

    def __init__(
        self,
        root: Path,
        files: Optional[Iterable[Path]] = None,
        known_hashes: Optional[Dict[str, str]] = None,
    ):
        """
        Initialize the analyzer with a project root and optional prior hashes.
        Args: root (Path): Project root to scan for documents.
              files (Optional[Iterable[Path]]): Optional iterable of files to analyze.
              known_hashes (Optional[Dict[str, str]]): Existing hash→path map to flag duplicates.
        Returns: None
        """
        self.root = Path(root)
        self.files = list(files) if files is not None else None
        self.known_hashes: Dict[str, str] = dict(known_hashes or {})

    def analyze(self) -> Dict[str, Any]:
        """
        Analyze supported documents under the root path and return structured findings.
        Args: None
        Returns: Dict[str, Any]: Contains per-document details, duplicates, hash index,
                 summary stats, and any errors encountered.
        """
        documents: List[Dict[str, Any]] = []
        duplicates: List[Dict[str, str]] = []
        errors: List[str] = []
        if self.files is None and not self.root.exists():
            return {
                "documents": [],
                "duplicates": [],
                "summary": {"unique_documents": 0, "duplicate_documents": 0, "total_words": 0, "by_format": {}},
                "hash_index": {},
                "errors": [f"Root path not found: {self.root}"],
            }
        paths = self.files if self.files is not None else sorted(self.root.rglob("*"))
        for path in paths:
            if not path.is_file():
                continue
            if path.name.startswith("._") or "__MACOSX" in path.parts:
                continue
            if ".git" in path.parts:
                continue
            suffix = path.suffix.lower()
            if suffix not in SUPPORTED_DOC_EXTS:
                continue
            try:
                rel_path = str(path.relative_to(self.root))
            except ValueError:
                rel_path = str(path)
            try:
                file_hash = compute_sha256(path)
            except Exception as e:
                errors.append(f"hash_failed:{rel_path}:{e}")
                continue
            if file_hash in self.known_hashes:
                duplicates.append(
                    {"path": rel_path, "hash": file_hash, "duplicate_of": self.known_hashes[file_hash]}
                )
                continue
            try:
                parsed = self._extract_content(path, suffix)
            except Exception as e:
                errors.append(f"parse_failed:{rel_path}:{e}")
                continue
            record = self._build_record(rel_path, file_hash, suffix, parsed)
            documents.append(record)
            self.known_hashes[file_hash] = rel_path
        summary = self._build_summary(documents, duplicates)
        return {
            "documents": documents,
            "duplicates": duplicates,
            "summary": summary,
            "hash_index": self.known_hashes,
            "errors": errors,
        }

    def _dedupe_preserve_order(
        self,
        items: List[Any],
        key_fn: Callable[[Any], Any] | None = None,
        limit: int | None = None,
    ) -> List[Any]:
        """
        Deduplicate items while preserving order.
        Args: items (List[Any]): Items to dedupe.
              key_fn (Callable[[Any], Any] | None): Optional key function.
              limit (int | None): Optional max items to return.
        Returns: List[Any]: Deduped list.
        """
        seen = set()
        out: List[Any] = []
        for item in items:
            key = key_fn(item) if key_fn else item
            if key in seen:
                continue
            seen.add(key)
            out.append(item)
            if limit is not None and len(out) >= limit:
                break
        return out

    def _build_summary(
        self,
        documents: List[Dict[str, Any]],
        duplicates: List[Dict[str, str]],
    ) -> Dict[str, Any]:
        """
        Summarize document counts and word totals for the analysis.
        Args: documents (List[Dict[str, Any]]): Unique document records.
              duplicates (List[Dict[str, str]]): Duplicate document entries.
        Returns: Dict[str, Any]: Aggregate counts for unique/duplicate docs, words, and formats.
        """
        total_words = sum(doc.get("word_count", 0) for doc in documents)
        by_format: Dict[str, int] = {}
        by_type: Dict[str, int] = {}
        for doc in documents:
            fmt = doc.get("format")
            if fmt:
                by_format[fmt] = by_format.get(fmt, 0) + 1
            doc_type = (doc.get("doc_type") or {}).get("label")
            if doc_type:
                by_type[doc_type] = by_type.get(doc_type, 0) + 1
        return {
            "unique_documents": len(documents),
            "duplicate_documents": len(duplicates),
            "total_words": total_words,
            "by_format": by_format,
            "by_type": by_type,
        }

    def _build_record(
        self,
        rel_path: str,
        file_hash: str,
        suffix: str,
        parsed: ParsedDoc,
    ) -> Dict[str, Any]:
        """
        Build the per-document record with extracted signals.
        Args: rel_path (str): File path relative to root.
              file_hash (str): SHA256 fingerprint of the file.
              suffix (str): File extension.
              parsed (ParsedDoc): Extracted text and headings.
        Returns: Dict[str, Any]: Structured document payload with metadata and heuristics.
        """
        text = parsed.text.strip()
        preview = " ".join(text.split())[:400]
        headings = parsed.headings[:10]
        lower_text = text.lower()
        title = self._extract_title(text, headings)
        venue = self._extract_venue(text)
        return {
            "path": rel_path,
            "format": suffix.lstrip(".").upper(),
            "sha256": file_hash,
            "word_count": self._word_count(text),
            "headings": headings,
            "preview": preview,
            "metrics": self._extract_metrics(text),
            "dates": self._extract_dates(text),
            "roles": self._extract_roles(lower_text, headings),
            "skills": self._extract_skills(lower_text),
            "doc_type": self._infer_doc_type(lower_text, headings),
            "topics": self._extract_topics(lower_text),
            "title": title,
            "summary": self._extract_summary(text, headings),
            "key_points": self._extract_key_points(text, headings),
            "authors": self._extract_authors(text, title),
            "venue": venue,
            "published_year": self._extract_year_from_venue(venue),
            "references_count": self._count_references(text),
            "figure_count": self._count_figures(text),
            "table_count": self._count_tables(text),
            "page_count": parsed.page_count,
        }

    def _extract_content(self, path: Path, suffix: str) -> ParsedDoc:
        """
        Route file reading based on extension.
        Args: path (Path): File path to read.
              suffix (str): Lowercased file extension.
        Returns: ParsedDoc: Extracted text and headings.
        """
        suffix = suffix.lower()
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix == ".md":
            return self._read_markdown(path)
        return self._read_plain_text(path)

    def _read_docx(self, path: Path) -> ParsedDoc:
        """
        Read DOCX content and collect paragraphs, headings, and table text.
        Args: path (Path): DOCX file path.
        Returns: ParsedDoc: Combined text and detected headings.
        """
        if Document is None: raise ImportError("python-docx is not installed")
        doc = Document(path)
        paragraphs = []
        headings = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text: continue
            paragraphs.append(text)
            try:
                style_name = (p.style.name or "").lower()
                if "heading" in style_name:
                    headings.append(text)
            except Exception:
                pass
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return ParsedDoc(text="\n".join(paragraphs), headings=headings)

    def _read_pdf(self, path: Path) -> ParsedDoc:
        """
        Read PDF content using pypdf and join page text.
        Args: path (Path): PDF file path.
        Returns: ParsedDoc: Combined text from all pages.
        """
        if PdfReader is None: raise ImportError("pypdf is not installed")
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt:
                pages.append(txt)
        return ParsedDoc(text="\n".join(pages), page_count=len(reader.pages))

    def _read_plain_text(self, path: Path) -> ParsedDoc:
        """
        Read a plain text file with utf-8 fallback to latin-1 on decode errors.
        Args: path (Path): Text file path.
        Returns: ParsedDoc: Raw text content.
        """
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
        return ParsedDoc(text=text)

    def _read_markdown(self, path: Path) -> ParsedDoc:
        """
        Read markdown files and capture heading lines.
        Args: path (Path): Markdown file path.
        Returns: ParsedDoc: Text content with detected headings.
        """
        parsed = self._read_plain_text(path)
        headings = []
        for line in parsed.text.splitlines():
            match = re.match(r"^#{1,6}\s+(.*)$", line.strip())
            if match:
                headings.append(match.group(1).strip())
        parsed.headings = headings
        return parsed

    def _word_count(self, text: str) -> int:
        """
        Count words in a block of text.
        Args: text (str): Text to count.
        Returns: int: Number of word tokens detected.
        """
        return len(re.findall(r"\b\w+\b", text))

    def _extract_metrics(self, text: str) -> List[str]:
        """
        Find metric-like patterns such as percentages and counts.
        Args: text (str): Text to search.
        Returns: List[str]: Unique metric strings found.
        """
        patterns = [
            r"\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*%",
            r"\b\d+(?:\.\d+)?\s*(?:users|clients|tickets|tests|deployments|issues|prs|pull requests)\b",
            r"\b\d+(?:\.\d+)?\s*(?:x|X|times)\b",
        ]
        metrics: List[str] = []
        for pat in patterns:
            metrics.extend(re.findall(pat, text, flags=re.IGNORECASE))
        metrics = [m.strip() for m in metrics]
        return self._dedupe_preserve_order(metrics, key_fn=lambda m: m.lower(), limit=10)

    def _extract_dates(self, text: str) -> List[str]:
        """
        Detect date-like ranges and years in text.
        Args: text (str): Text to search.
        Returns: List[str]: Unique date strings found.
        """
        patterns = [
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*(?:-|to|through|–|—)\s*(?:Present|Now|\d{4})",
            r"\b\d{4}\s*(?:-|to|through|–|—)\s*(?:Present|Now|\d{4})",
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
            r"\b(20\d{2}|19\d{2})\b",
        ]
        context_keywords = {"published", "accepted", "conference", "journal", "arxiv", "doi", "copyright", "submitted", "proceedings", "volume", "issue"}
        dates: List[str] = []
        header_text = text[:2000]
        for pat in patterns:
            dates.extend(re.findall(pat, header_text, flags=re.IGNORECASE))
        filtered: List[str] = []
        for d in dates:
            val = d if isinstance(d, str) else str(d)
            if re.search(r"\b\d{3,4}\s*[–—-]\s*\d{3,4}\b", val): continue
            if re.fullmatch(r"(20\d{2}|19\d{2})", val):
                idx = header_text.lower().find(val.lower())
                window = header_text.lower()[max(0, idx - 40): idx + 40] if idx >= 0 else ""
                if not any(k in window for k in context_keywords):
                    continue
            filtered.append(val)
        return self._dedupe_preserve_order(filtered, key_fn=lambda v: v.lower(), limit=10)

    def _extract_roles(self, lower_text: str, headings: List[str]) -> List[str]:
        """
        Identify role-related keywords in text and headings.
        Args: lower_text (str): Lowercased document text.
              headings (List[str]): Extracted headings.
        Returns: List[str]: Sorted list of role titles found.
        """
        role_keywords = ["engineer", "developer", "manager", "lead", "researcher", "analyst", "designer", "architect", "consultant"]
        roles = set()
        text_slice = lower_text[:1200]
        for keyword in role_keywords:
            if keyword in text_slice:
                roles.add(keyword.title())
        for heading in headings:
            h_lower = heading.lower()
            for keyword in role_keywords:
                if keyword in h_lower:
                    roles.add(keyword.title())
        return sorted(roles)

    def _infer_doc_type(self, lower_text: str, headings: List[str]) -> Dict[str, Any]:
        """
        Heuristically infer a document type from keywords and headings.
        Args: lower_text (str): Lowercased document text.
              headings (List[str]): Extracted headings.
        Returns: Dict[str, Any]: label, confidence, and matched signals.
        """
        heading_blob = " ".join(h.lower() for h in headings)
        rules = [
            ("research paper", ["abstract", "introduction", "method", "results", "conclusion", "references"]),
            ("resume/cv", ["resume", "curriculum vitae", "experience", "education", "skills", "objective", "summary"]),
            ("proposal", ["proposal", "scope of work", "statement of work", "sow", "deliverables", "timeline"]),
            ("report", ["report", "findings", "analysis", "results", "summary", "conclusion"]),
            ("research", ["research", "methodology", "hypothesis", "literature review", "dataset", "experiment"]),
            ("policy", ["policy", "compliance", "regulation", "governance", "risk management"]),
            ("specification", ["specification", "requirements", "acceptance criteria", "user story", "use case"]),
            ("manual/guide", ["guide", "manual", "instructions", "how to", "procedure", "step-by-step"]),
            ("minutes/agenda", ["meeting minutes", "minutes", "agenda", "action items", "attendees"]),
            ("invoice/receipt", ["invoice", "receipt", "bill to", "amount due", "total due"]),
            ("budget/financial", ["budget", "forecast", "revenue", "expense", "balance sheet", "cash flow"]),
            ("presentation", ["slide", "deck", "presentation", "speaker notes"]),
            ("software/technical", ["api", "endpoint", "code", "repository", "build", "deployment"]),
        ]
        best_label = "unknown"
        best_score = 0
        matched_signals: List[str] = []
        for label, keywords in rules:
            score = 0
            local_hits: List[str] = []
            for kw in keywords:
                pattern = r"\b" + re.escape(kw) + r"\b"
                in_text = re.search(pattern, lower_text) is not None
                in_headings = re.search(pattern, heading_blob) is not None
                if in_text:
                    score += 1
                    local_hits.append(kw)
                if in_headings:
                    score += 2
                    local_hits.append(kw)
            if score > best_score:
                best_score = score
                best_label = label
                matched_signals = local_hits
        if best_label == "resume/cv" and best_score < 3:
            best_label = "unknown"
            best_score = 0
            matched_signals = []
        if best_score >= 4:
            confidence = "high"
        elif best_score >= 2:
            confidence = "medium"
        elif best_score >= 1:
            confidence = "low"
        else:
            confidence = "unknown"
        unique_signals = self._dedupe_preserve_order(matched_signals, limit=8)
        return {"label": best_label, "confidence": confidence, "signals": unique_signals}

    def _extract_skills(self, lower_text: str) -> List[str]:
        """
        Identify predefined skills from document text.
        Args: lower_text (str): Lowercased document text.
        Returns: List[str]: Ordered unique skills detected.
        """
        skill_terms = {
            "python": "Python", "java": "Java", "javascript": "JavaScript", "typescript": "TypeScript", "c++": "C++", "c+": "C++",
            "cpp": "C++", "c#": "C#", "golang": "Go", "go language": "Go", "rust": "Rust", "sql": "SQL",
            "react": "React", "node.js": "Node.js", "nodejs": "Node.js", "express": "Express", "django": "Django", "flask": "Flask",
            "fastapi": "FastAPI", "spring": "Spring", "angular": "Angular", "vue": "Vue", "next.js": "Next.js", "nestjs": "NestJS",
            "pytorch": "PyTorch", "tensorflow": "TensorFlow", "docker": "Docker",
        }
        detected = []
        for needle, label in skill_terms.items():
            if needle in lower_text:
                detected.append(label)
        return self._dedupe_preserve_order(detected)

    def _extract_topics(self, lower_text: str) -> List[str]:
        """
        Extract top non-trivial keywords as topic hints.
        Args: lower_text (str): Lowercased document text.
        Returns: List[str]: Top topic keywords.
        """
        stopwords = {
            "the", "and", "for", "with", "from", "that",
            "this", "these", "those", "your", "you", "our",
            "are", "was", "were", "have", "has", "had",
            "not", "but", "into", "over", "under", "between",
            "within", "about", "project", "document", "report", "analysis",
            "summary", "section", "page", "pages", "file", "files",
            "data", "information", "results", "more", "less", "than",
            "also", "use", "using", "used", "arxiv", "paper",
            "et", "al", "figure", "table",
        }
        words = re.findall(r"\b[a-z][a-z0-9\-]{3,}\b", lower_text)
        freq: Dict[str, int] = {}
        for w in words:
            if w in stopwords: continue
            freq[w] = freq.get(w, 0) + 1
        ranked = sorted(freq.items(), key=lambda kv: (-kv[1], kv[0]))
        return [w for w, _ in ranked[:12]]

    def _extract_title(self, text: str, headings: List[str]) -> str:
        """
        Extract a display title from headings or the first meaningful line.
        Args: text (str): Full document text.
              headings (List[str]): Extracted headings.
        Returns: str: Title candidate.
        """
        if headings:
            return headings[0][:120]
        ignore_snippets = ["provided proper attribution", "permission to reproduce", "copyright", "all rights reserved", "arxiv", "doi:", "conference", "proceedings", "submitted", "accepted", "equal contribution", "listing order", "corresponding author", "http://", "https://"]
        verb_snippets = ["proposed", "started", "developed", "evaluated", "presented"]
        best = ""
        best_score = 0
        def score_line(line: str) -> int:
            words = re.findall(r"[A-Za-z][A-Za-z\-']+", line)
            if not words: return -999
            if len(words) < 3 or len(words) > 15: return -999
            title_case = sum(1 for w in words if w[:1].isupper())
            ratio = title_case / max(1, len(words))
            if ratio < 0.6: return -999
            if re.search(r"\d|%", line): return -999
            long_words = sum(1 for w in words if len(w) >= 4)
            score = (title_case * 2) + long_words - abs(len(words) - 8)
            if ":" in line or "—" in line or "-" in line:
                score += 1
            return score
        lower_all = text.lower()
        abstract_idx = lower_all.find("abstract")
        if abstract_idx != -1:
            pre = text[max(0, abstract_idx - 800):abstract_idx]
            for raw in pre.splitlines():
                line = raw.strip()
                if not line:
                    continue
                lower = line.lower()
                if any(snippet in lower for snippet in ignore_snippets):
                    continue
                if any(v in lower for v in verb_snippets):
                    continue
                score = score_line(line)
                if score > best_score:
                    best = line
                    best_score = score
        for raw in text.splitlines()[:120]:
            line = raw.strip()
            if not line: continue
            if len(line) < 8 or len(line) > 140: continue
            lower = line.lower()
            if any(snippet in lower for snippet in ignore_snippets): continue
            if any(v in lower for v in verb_snippets): continue
            if re.search(r"^\d+\s*$", line): continue
            score = score_line(line)
            if score > best_score:
                best = line
                best_score = score
        if best:
            return best
        for line in text.splitlines():
            line = line.strip()
            if not line: continue
            if len(line) < 6 or len(line) > 160: continue
            return line
        return "Untitled document"

    def _extract_summary(self, text: str, headings: List[str]) -> str:
        """
        Extract a short summary (1-2 sentences) using simple scoring.
        Args: text (str): Full document text.
              headings (List[str]): Extracted headings.
        Returns: str: Summary sentence(s).
        """
        cleaned = self._clean_text_for_summary(text)
        abstract = self._extract_section(cleaned, headings, "abstract")
        if not abstract:
            abstract = self._extract_section(cleaned, headings, "abstract", end_kw="introduction", prefer_headings=False)
        if abstract:
            return self._summarize_block(abstract)
        sentences = re.split(r"(?<=[.!?])\s+", cleaned.strip())
        scored = []
        keywords = {"goal", "objective", "purpose", "summary", "overview", "finding", "results", "conclusion"}
        for idx, s in enumerate(sentences):
            s_clean = " ".join(s.split())
            if len(s_clean) < 40: continue
            word_count = len(re.findall(r"\b\w+\b", s_clean))
            if word_count < 8 or word_count > 40: continue
            lower = s_clean.lower()
            score = 0
            if idx < 3:
                score += 2
            if any(k in lower for k in keywords):
                score += 2
            if re.search(r"\b\d{4}\b", s_clean):
                score += 1
            if re.search(r"\b\d+(?:\.\d+)?\s*%\b", s_clean):
                score += 1
            scored.append((score, idx, s_clean))
        if not scored:
            return ""
        top = sorted(scored, key=lambda t: (-t[0], t[1]))[:2]
        top_sorted = sorted(top, key=lambda t: t[1])
        return " ".join(s for _, _, s in top_sorted)[:420]

    def _summarize_block(self, text: str) -> str:
        """
        Summarize a block of text using a couple of strong sentences.
        Args: text (str): Text block to summarize.
        Returns: str: Short summary text.
        """
        sentences = re.split(r"(?<=[.!?])\s+", text.strip())
        picked = []
        for s in sentences:
            s_clean = " ".join(s.split())
            if 50 <= len(s_clean) <= 220:
                picked.append(s_clean)
            if len(picked) >= 2: break
        return " ".join(picked)[:420]

    def _extract_section(self, text: str, headings: List[str], name: str, end_kw: str | None = None, prefer_headings: bool = True) -> str:
        """
        Extract a section body based on heading match or keyword markers.
        Args: text (str): Full document text.
              headings (List[str]): Extracted headings.
              name (str): Section name to extract.
              end_kw (str | None): Optional end marker when keyword scanning.
              prefer_headings (bool): Prefer heading-based extraction when available.
        Returns: str: Extracted section text.
        """
        lines = text.splitlines()
        if prefer_headings and not headings:
            return ""
        start_idx = None
        end_idx = None
        for i, line in enumerate(lines):
            if re.match(rf"^\s*{re.escape(name)}\b", line.strip(), flags=re.IGNORECASE):
                start_idx = i + 1
                break
        if start_idx is None and not prefer_headings:
            lower = text.lower()
            start_idx = lower.find(name.lower())
            if start_idx == -1: return ""
            end_idx = lower.find((end_kw or "").lower(), start_idx + len(name)) if end_kw else -1
            if end_idx == -1:
                end_idx = start_idx + 1200
            chunk = text[start_idx:end_idx].strip()
            chunk = re.sub(rf"^\s*{re.escape(name)}\s*[:\-]?\s*", "", chunk, flags=re.IGNORECASE)
            return chunk[:1200]
        if start_idx is None:
            return ""
        for j in range(start_idx, len(lines)):
            if re.match(r"^\s*[A-Z][A-Za-z0-9\s\-]{2,}\s*$", lines[j].strip()):
                end_idx = j
                break
        chunk = "\n".join(lines[start_idx:end_idx]).strip()
        return chunk[:1200]

    def _clean_text_for_summary(self, text: str) -> str:
        """
        Remove boilerplate lines that pollute summaries.
        Args: text (str): Full document text.
        Returns: str: Cleaned text.
        """
        drop_snippets = ["provided proper attribution", "permission to reproduce", "reproduce the tables and figures", "all rights reserved", "copyright", "arxiv", "doi:", "conference on", "proceedings"]
        cleaned_lines = []
        for line in text.splitlines():
            l = line.strip()
            if not l: continue
            lower = l.lower()
            if any(s in lower for s in drop_snippets): continue
            cleaned_lines.append(l)
        return "\n".join(cleaned_lines)

    def _extract_key_points(self, text: str, headings: List[str]) -> List[str]:
        """
        Extract key points from bullet-like lines or salient sentences.
        Args: text (str): Full document text.
              headings (List[str]): Extracted headings.
        Returns: List[str]: Short key points.
        """
        conclusion = self._extract_section(text, headings, "conclusion")
        if not conclusion:
            conclusion = self._extract_section(text, headings, "conclusion", end_kw="references", prefer_headings=False)
        if conclusion:
            points = self._summarize_block(conclusion)
            if points: return [points]
        points: List[str] = []
        for line in text.splitlines():
            l = line.strip()
            if not l: continue
            if re.match(r"^[-*•]\s+", l):
                points.append(re.sub(r"^[-*•]\s+", "", l))
            if len(points) >= 6: break
        if points:
            return points[:6]
        sentences = re.split(r"(?<=[.!?])\s+", text.strip())
        verbs = {"built", "created", "designed", "implemented", "analyzed", "evaluated", "proposed", "authored", "drafted"}
        scored = []
        for s in sentences:
            s_clean = " ".join(s.split())
            if len(s_clean) < 50: continue
            lower = s_clean.lower()
            score = 0
            if any(v in lower for v in verbs):
                score += 2
            if re.search(r"\b\d{4}\b", s_clean):
                score += 1
            if re.search(r"\b\d+(?:\.\d+)?\s*%\b", s_clean):
                score += 1
            if score > 0:
                scored.append((score, s_clean))
        scored = sorted(scored, key=lambda t: (-t[0], len(t[1])))
        return [s for _, s in scored[:1]]

    def _extract_year_from_venue(self, venue: str) -> str:
        """
        Pull a year out of the venue line if present.
        Args: venue (str): Venue line.
        Returns: str: Year if found, else empty string.
        """
        if not venue:
            return ""
        match = re.search(r"\b(19\d{2}|20\d{2})\b", venue)
        return match.group(1) if match else ""

    def _extract_authors(self, text: str, title: str) -> List[str]:
        """
        Extract author-like names near the top of the document.
        Args: text (str): Full document text.
              title (str): Title text to exclude from author candidates.
        Returns: List[str]: Author name candidates.
        """
        lines = [l.strip() for l in text.splitlines()[:80] if l.strip()]
        ignore_snippets = ["provided proper attribution", "permission to reproduce", "copyright", "all rights reserved", "conference", "proceedings", "arxiv"]
        author_lines = []
        for line in lines:
            if "abstract" in line.lower():
                break
            lower = line.lower()
            if any(s in lower for s in ignore_snippets):
                continue
            if "," in line and re.search(r"[A-Z][a-z]+", line):
                author_lines.append(line)
        joined = " ".join(author_lines)
        org_terms = {"University", "Institute", "Laboratory", "Labs", "Research", "Department", "School", "College", "Inc", "Ltd", "LLC", "Company", "Google", "Microsoft", "Facebook", "Meta", "Amazon", "Apple"}
        blacklist_words = {"Abstract", "Conference", "Proceedings", "Information", "Systems", "Beach", "Attention", "Neural", "Need", "The", "Long", "Beach", "Introduction", "Recurrent", "Convolutional", "Decoder", "Encoder"}
        title_words = {w for w in re.findall(r"[A-Za-z]+", title) if len(w) > 3}
        if not joined:
            header = " ".join(lines[:120])
            joined = header
        candidates = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}\b", joined)
        authors = []
        for name in candidates:
            if any(term in name for term in org_terms):
                continue
            if any(word in name.split() for word in blacklist_words):
                continue
            if any(word in name for word in title_words):
                continue
            if len(name) > 60:
                continue
            authors.append(name)
            if len(authors) >= 10:
                break
        return self._dedupe_preserve_order(authors)

    def _extract_venue(self, text: str) -> str:
        """
        Extract a venue/conference line if present.
        Args: text (str): Full document text.
        Returns: str: Venue line if found.
        """
        for line in text.splitlines()[:120]:
            l = line.strip()
            if not l:
                continue
            if re.search(r"\bconference\b|\bproceedings\b|\bjournal\b|\bworkshop\b", l, re.IGNORECASE):
                if len(l) <= 160:
                    return l
        return ""

    def _count_references(self, text: str) -> int:
        """
        Count reference-like entries from a references section.
        Args: text (str): Full document text.
        Returns: int: Approximate reference count.
        """
        lower = text.lower()
        idx = lower.find("references")
        if idx == -1:
            return 0
        tail = text[idx: idx + 3000]
        hits = re.findall(r"\n\s*\[\d+\]\s+", tail)
        return min(len(hits), 200)

    def _count_figures(self, text: str) -> int:
        """
        Count figure mentions.
        Args: text (str): Full document text.
        Returns: int: Number of figure mentions.
        """
        return min(len(re.findall(r"\bfig(?:ure)?\.?\s*\d+", text, flags=re.IGNORECASE)), 200)

    def _count_tables(self, text: str) -> int:
        """
        Count table mentions.
        Args: text (str): Full document text.
        Returns: int: Number of table mentions.
        """
        return min(len(re.findall(r"\btable\s*\d+", text, flags=re.IGNORECASE)), 200)
