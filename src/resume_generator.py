# Builds Word resumes for projects and contributors

import os
from datetime import datetime
import ctypes

from docx import Document
from docx.shared import Pt
from file_parser import OUTPUT_DIR
from db import update_full_scan, list_full_scans, get_full_scan_by_id
from print_utils import _center_text, is_noise
from permission_manager import get_yes_no


RESUME_DIR = os.path.join(OUTPUT_DIR, "resumes")
os.makedirs(RESUME_DIR, exist_ok=True)

# -------------------------------------------------------------------------
# SHARED HELPERS
# -------------------------------------------------------------------------

def _fmt_date(val):
    """Formats a datetime or ISO string into YYYY-MM-DD."""
    if isinstance(val, datetime):
        return val.date().isoformat()
    if isinstance(val, str):
        return val.split("T")[0]
    return ""


def _save_doc(doc, path):
    """Helper to save a docx with retry logic."""
    while True:
        try:
            doc.save(path)
            return path
        except PermissionError:
            print(f"\n[!] Could not save to '{path}' because it is open.")
            print("Please close the file and press Enter to retry, or type 'cancel' to stop.")
            if input("Action: ").strip().lower() == "cancel":
                return None
        except Exception as e:
            print(f"Error saving document: {e}")
            return None


def _input_with_prefill(prompt, text):
    """
    Prompts for input with a default value pre-filled.
    Works on Windows (via ctypes) and Unix (via readline).
    """
    if not text:
        return input(prompt)

    if os.name == 'nt':
        try:
            # Windows implementation using WriteConsoleInputW
            class KEY_EVENT_RECORD(ctypes.Structure):
                _fields_ = [
                    ("bKeyDown", ctypes.c_int),
                    ("wRepeatCount", ctypes.c_ushort),
                    ("wVirtualKeyCode", ctypes.c_ushort),
                    ("wVirtualScanCode", ctypes.c_ushort),
                    ("uChar", ctypes.c_wchar),
                    ("dwControlKeyState", ctypes.c_ulong)
                ]
            
            class INPUT_RECORD_Event(ctypes.Union):
                _fields_ = [("KeyEvent", KEY_EVENT_RECORD)]
            
            class INPUT_RECORD(ctypes.Structure):
                _fields_ = [
                    ("EventType", ctypes.c_ushort),
                    ("Event", INPUT_RECORD_Event)
                ]

            STD_INPUT_HANDLE = -10
            hConsoleInput = ctypes.windll.kernel32.GetStdHandle(STD_INPUT_HANDLE)
            
            n = len(text)
            records = (INPUT_RECORD * n)()
            
            for i, char in enumerate(text):
                records[i].EventType = 0x0001 # KEY_EVENT
                records[i].Event.KeyEvent.bKeyDown = 1
                records[i].Event.KeyEvent.wRepeatCount = 1
                records[i].Event.KeyEvent.wVirtualKeyCode = 0
                records[i].Event.KeyEvent.wVirtualScanCode = 0
                records[i].Event.KeyEvent.uChar = char
                records[i].Event.KeyEvent.dwControlKeyState = 0

            written = ctypes.c_ulong(0)
            ctypes.windll.kernel32.WriteConsoleInputW(
                hConsoleInput,
                ctypes.byref(records),
                n,
                ctypes.byref(written)
            )
        except Exception:
            pass
    else:
        try:
            import readline
            def hook():
                readline.insert_text(text)
                readline.redisplay()
            readline.set_startup_hook(hook)
        except ImportError:
            pass

    try:
        return input(prompt)
    finally:
        if os.name != 'nt':
            try:
                import readline
                readline.set_startup_hook()
            except ImportError:
                pass


# -------------------------------------------------------------------------
# 1. GENERAL PROJECT RESUME (Summary of all projects)
# -------------------------------------------------------------------------


def build_project_line(p: dict) -> str:
    """
    Builds a short resume description about a project (General/Team view).
    Used for the full resume summary.
    """
    name = p.get("project", "Unknown")
    langs = p.get("languages", "Unknown")
    skills = p.get("skills", "NA")
    frameworks = p.get("frameworks", "None")
    duration = p.get("duration_days", 0)
    code_files = p.get("code_files", 0)
    test_files = p.get("test_files", 0)
    project_type = p.get("project_type", "software")

    # Main clause
    main = f"Contributed to project '{name}'"
    if project_type and project_type.lower() != "unknown":
        main = f"Contributed to {project_type.lower()} project '{name}'"
    if langs and langs.lower() != "unknown":
        main += f" using {langs}"

    # Details clause
    details = []
    if code_files:
        details.append(f"{code_files} code files")
    if test_files:
        details.append(f"{test_files} test files")

    if duration:
        details.append(f"over {duration} days")
    
    pieces = [main]
    if details:
        pieces.append("comprising " + ", ".join(details))

    if frameworks and frameworks not in ("None", "NA"):
        pieces.append(f"with frameworks such as {frameworks}")

    return "; ".join(pieces) + "."


def generate_resume(
    project_summaries: list[dict],
    chronological_projects: list[dict],
    skills_output: list[dict],
    scan_timestamp: str = None
) -> str:
    """
    Generates the 'Full Project Resume' Word document.
    Returns the path to the generated file.
    """
    # Sort projects by score
    sorted_projects = sorted(
        project_summaries, key=lambda p: p.get("score", 0), reverse=True
    )

    docx_path = os.path.join(RESUME_DIR, "portfolio_resume.docx")

    doc = Document()

    # Title
    title = doc.add_heading("Project Portfolio Resume", level=0)
    title.runs[0].font.size = Pt(20)

    # Placeholder Info
    info_p = doc.add_paragraph()
    info_p.add_run("Generated by Skill Scope").italic = True
    doc.add_paragraph()

    # Top Projects
    doc.add_heading("Top Projects", level=1)
    if not sorted_projects:
        doc.add_paragraph("No projects detected.", style="List Bullet")
    else:
        for p in sorted_projects:
            first = _fmt_date(p.get("first_modified"))
            last = _fmt_date(p.get("last_modified"))

            para = doc.add_paragraph(style="List Bullet")
            para.add_run(p["project"]).bold = True
            if first and last:
                para.add_run(f"  ({first} – {last})")
            para.add_run("\n" + build_project_line(p))

    doc.add_paragraph()

    # Timeline
    doc.add_heading("Project Timeline", level=1)
    for p in chronological_projects:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(p["name"]).bold = True
        para.add_run(f" – {p['first_used']} → {p['last_used']}")
    doc.add_paragraph()

    # Skills
    doc.add_heading("Skills Used Over Time", level=1)
    if skills_output:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Skill"
        hdr[1].text = "First Used"
        hdr[2].text = "Last Used"
        for row in skills_output:
            cells = table.add_row().cells
            cells[0].text = row["skill"]
            cells[1].text = row["first_used"]
            cells[2].text = row["last_used"]
    else:
        doc.add_paragraph("No skills detected.", style="List Bullet")

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph("Generated automatically from code repositories using Skill Scope.")
    foot.runs[0].italic = True
    foot.runs[0].font.size = Pt(8)

    return _save_doc(doc, docx_path)


# -------------------------------------------------------------------------
# 2. INDIVIDUAL CONTRIBUTOR RESUME (Specific Person)
# -------------------------------------------------------------------------


def render_resume_artifact(artifact_data: dict, output_path: str) -> str:
    """
    Renders a DOCX file from a resume artifact dictionary.
    This decouples the data structure from the file generation, allowing
    both the CLI and API to use the same rendering logic.
    """
    doc = Document()
    
    # 1. Header Info
    user_name = artifact_data.get("user_name") or "Resume"
    user_title = artifact_data.get("user_title")
    user_summary = artifact_data.get("user_summary")
    created_date = artifact_data.get("created_date")
    
    # Title / Name
    h1 = doc.add_heading(user_name, level=0)
    h1.runs[0].font.size = Pt(24)
    
    if created_date:
        doc.add_paragraph(f"Generated on {created_date}")
    
    if user_title:
        p = doc.add_paragraph(user_title)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph() # Spacer

    # 2. Summary Section
    if user_summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(user_summary)
        doc.add_paragraph()

    # 3. Skills Section
    skills = artifact_data.get("skills", [])
    if skills:
        doc.add_heading("Technical Skills", level=1)
        p = doc.add_paragraph()
        p.add_run("Languages & Technologies: ").bold = True
        p.add_run(", ".join(skills))
        doc.add_paragraph()

    # 4. Projects Section
    items = artifact_data.get("items", [])
    if items:
        doc.add_heading("Project Experience", level=1)
        
        for item in items:
            project_name = item.get("project_name", "Unnamed Project")
            text = item.get("text", "")
            date_str = item.get("date_str", "")
            
            # Project Header
            p_head = doc.add_heading(level=2)
            p_head.add_run(project_name).bold = True
            if date_str:
                p_head.add_run(date_str).font.size = Pt(11)
            
            # Description
            if text:
                doc.add_paragraph(text, style="List Bullet")
            
            # Per-project skills (optional, if present in item)
            project_skills = item.get("skills", [])
            if project_skills:
                s_p = doc.add_paragraph(style="List Bullet")
                s_p.add_run("Skills: ").bold = True
                s_p.add_run(", ".join(project_skills))
            
            doc.add_paragraph()

    return _save_doc(doc, output_path)


def _build_personal_project_description(project_name, project_context, user_stats):
    """
    Constructs a sentence describing the user's specific contribution to a project.
    """
    # Context from the project as a whole
    langs = project_context.get("languages", "Unknown")
    skills = project_context.get("skills", "NA")
    frameworks = project_context.get("frameworks", "None")
    
    # User specific stats
    u_files = user_stats.get("files_worked", 0)
    u_code = user_stats.get("user_code_files", 0)
    u_test = user_stats.get("user_test_files", 0)
    u_doc = user_stats.get("user_doc_files", 0)
    u_design = user_stats.get("user_design_files", 0)
    
    # Determine the most appropriate action verb based on work distribution
    verb = "Contributed to"
    total_work = u_code + u_test + u_doc + u_design
    
    if total_work > 0:
        if u_code >= u_test and u_code >= u_doc and u_code >= u_design:
            verb = "Developed key components for"
        elif u_test > u_code and u_test > u_doc:
            verb = "Implemented testing suites for"
        elif u_doc > u_code and u_doc > u_test:
            verb = "Authored technical documentation for"
        elif u_design > u_code:
            verb = "Designed assets and UI elements for"

    parts = []
    parts.append(f"{verb} project development")
    
    # Quantifiable Metrics (Suggestion 5)
    pct = user_stats.get("pct", 0.0)
    if pct > 10.0:
        parts.append(f", contributing {pct:.1f}% of the codebase")
        
    if u_files > 0:
        parts.append(f"impacting {u_files} files")

    duration = project_context.get("duration_days", 0)
    if duration > 14:
        parts.append(f"over a {duration}-day period")
    
    if langs and langs != "Unknown":
        parts.append(f"using {langs}")
    if frameworks and frameworks not in ("None", "NA"):
        parts.append(f"utilizing frameworks such as {frameworks}")
        
    return " ".join(parts) + "."


def generate_contributor_portfolio(
    contributor_name: str,
    profile_data: dict,
    all_projects_map: dict,
    scan_timestamp: str = None,
    sort_mode: str = "impact"
) -> str:
    """
    Generates a specific resume Word doc for a single contributor.
    """
    if not profile_data:
        return None

    # 1. Prepare Data Structure (The "Artifact")
    safe_name = "".join(c for c in contributor_name if c.isalnum() or c in (' ', '_', '-')).strip()
    docx_path = os.path.join(RESUME_DIR, f"Resume_{safe_name}.docx")
    
    artifact = {
        "created_date": datetime.now().strftime('%B %d, %Y'),
        "items": []
    }
    
    skills = profile_data.get("skills", [])
    projects_ref = profile_data.get("projects", [])
    
    # Determine a professional title based on skills
    role = profile_data.get("custom_title")
    if not role:
        role = "Software Contributor"
        dev_keywords = {"Development", "Programming", "Engineering"}
        if any(any(k in s for k in dev_keywords) for s in skills):
            role = "Software Developer"

    # Name
    custom_name = profile_data.get("custom_name")
    if custom_name:
        display_name = custom_name
    else:
        display_name = contributor_name
        if "@" in contributor_name:
            display_name = contributor_name.split("@")[0]
        display_name = display_name.replace(".", " ").replace("_", " ").title()
    
    artifact["user_name"] = display_name
    artifact["user_title"] = role
    artifact["skills"] = skills

    custom_summary = profile_data.get("custom_summary")
    if custom_summary:
        summary_text = custom_summary
    else:
        summary_text = f"{role} with a track record of contributions across {len(projects_ref)} project(s)."

        if skills:
            # List top 3 skills naturally
            top_skills = skills[:3]
            summary_text += f" Proficient in {', '.join(top_skills)}"
            if len(skills) > 3:
                summary_text += f", along with expertise in {len(skills)-3} other technologies"
            summary_text += "."
        
    artifact["user_summary"] = summary_text
    
    # Prepare projects
    user_projects = []
    for p_ref in projects_ref:
        p_name = p_ref.get("name", "Unknown Project")
        
        # Reconstruct stats
        user_stats = {
            "user_code_files": p_ref.get("user_code_files", 0),
            "user_test_files": p_ref.get("user_test_files", 0),
            "user_doc_files": p_ref.get("user_doc_files", 0),
            "user_design_files": p_ref.get("user_design_files", 0),
            "pct": p_ref.get("pct", 0.0),
            "score": p_ref.get("score", 0.0),
            "files_worked": p_ref.get("files_worked", 0),
            "commit_count": p_ref.get("commit_count", 0)
        }

        # Fallback for files_worked
        if user_stats["files_worked"] == 0 and p_ref.get("files_list"):
             user_stats["files_worked"] = len(p_ref.get("files_list"))

        project_context = all_projects_map.get(p_name, {})
        custom_desc = p_ref.get("custom_description")
        custom_skills = p_ref.get("custom_skills")
        user_projects.append((p_name, user_stats, project_context, custom_desc, custom_skills))

    # Sort by impact
    user_projects.sort(key=lambda x: x[1]["score"], reverse=True)
    # Sort Projects (Suggestion 6)
    if sort_mode == "date":
        # Chronological (Newest First)
        # Ensure we handle both datetime objects and ISO strings safely
        user_projects.sort(
            key=lambda x: str(x[2].get("last_modified", "")), reverse=True
        )
    else:
        # Functional (Impact Score - Best First)
        user_projects.sort(key=lambda x: x[1]["score"], reverse=True)

    for p_name, u_stats, p_context, custom_desc, custom_skills in user_projects:
        # Filter negligible
        if u_stats["pct"] < 0.1 and u_stats["files_worked"] == 0 and u_stats["commit_count"] == 0:
            continue

        # Date
        start = p_context.get("first_modified")
        end = p_context.get("last_modified")
        date_str = ""
        if start and end:
            date_str = f" ({_fmt_date(start)} – {_fmt_date(end)})"

        # Description
        if custom_desc:
            desc = custom_desc
        else:
            desc = _build_personal_project_description(p_name, p_context, u_stats)
        
        # Skills for this project
        if custom_skills is not None:
            my_skills = custom_skills
        else:
            pcs = p_context.get("per_contributor_skills", {})
            my_skills = pcs.get(contributor_name, [])

        artifact["items"].append({
            "project_name": p_name,
            "text": desc,
            "date_str": date_str,
            "skills": my_skills
        })

    # 2. Render Artifact
    return render_resume_artifact(artifact, docx_path)


def edit_contributor_descriptions(target_scan=None):
    """
    Allows the user to edit the description of a project for a specific contributor.
    """
    if target_scan:
        scan = target_scan
        summary_id = target_scan["summary_id"]
        data = target_scan["scan_data"]
    else:
        scans = list_full_scans()
        if not scans:
            print(_center_text("No scans found."))
            return

        print()
        print(_center_text("Select a scan to edit:"))
        for i, s in enumerate(scans, start=1):
            print(_center_text(f"{i}. Scan {s['summary_id']} - {s['timestamp']}"))

        choice = input(_center_text("Enter number (0 to cancel): ")).strip()
        if not choice.isdigit() or int(choice) == 0:
            return

        idx = int(choice) - 1
        if idx < 0 or idx >= len(scans):
            print(_center_text("Invalid selection."))
            return

        summary_id = scans[idx]["summary_id"]
        scan = get_full_scan_by_id(summary_id)
        if not scan:
            print(_center_text("Error loading scan data."))
            return
        data = scan["scan_data"]

    profiles = data.get("contributor_profiles", {})
    project_summaries = data.get("project_summaries", [])
    project_map = {p["project"]: p for p in project_summaries}

    contributors = sorted([c for c in profiles.keys() if not is_noise(c)])

    if not contributors:
        print(_center_text("No contributors found."))
        return

    while True:
        print()
        print(_center_text("Select contributor to edit:"))
        for i, c in enumerate(contributors, 1):
            print(_center_text(f"{i}. {c}"))

        sel = input(_center_text("Enter number (0 to back): ")).strip()
        if not sel.isdigit():
            continue
        
        c_idx = int(sel) - 1
        if c_idx == -1:
            break
        if c_idx < 0 or c_idx >= len(contributors):
            continue

        user = contributors[c_idx]
        profile = profiles[user]
        user_projects = profile.get("projects", [])
        skills = profile.get("skills", [])

        while True:
            print()
            print(_center_text(f"--- Editing {user} ---"))
            
            menu_opts = [
                "1. Edit Name",
                "2. Edit Professional Title",
                "3. Edit Professional Summary",
                "4. Edit Project Details (Desc/Skills)",
                "5. Regenerate Resume",
                "6. Reset Resume Changes",
                "0. Back to Contributor List"
            ]
            max_len = max(len(o) for o in menu_opts)
            for opt in menu_opts:
                print(_center_text(opt.ljust(max_len)))

            choice = input(_center_text("Choose option: ")).strip()

            if choice == "0":
                break

            if choice == "1":
                # Calculate default name
                default_name = user
                if "@" in user:
                    default_name = user.split("@")[0]
                default_name = default_name.replace(".", " ").replace("_", " ").title()

                curr = profile.get("custom_name", default_name)
                print(_center_text("Edit Name (type 'RESET' to restore default):"))
                val = _input_with_prefill("Value: ", curr).strip()

                if val == "RESET":
                    if "custom_name" in profile:
                        del profile["custom_name"]
                        update_full_scan(summary_id, data)
                        print(_center_text("Reset to default."))
                    else:
                        print(_center_text("Already default."))
                elif val:
                    profile["custom_name"] = val
                    update_full_scan(summary_id, data)
                    print(_center_text("Saved."))
                else:
                    print(_center_text("No change."))

            elif choice == "2":
                # Calculate default title
                default_title = "Software Contributor"
                dev_keywords = {"Development", "Programming", "Engineering"}
                if any(any(k in s for k in dev_keywords) for s in skills):
                    default_title = "Software Developer"

                curr = profile.get("custom_title", default_title)
                print(_center_text("Edit Title (type 'RESET' to restore default):"))
                val = _input_with_prefill("Value: ", curr).strip()

                if val == "RESET":
                    if "custom_title" in profile:
                        del profile["custom_title"]
                        update_full_scan(summary_id, data)
                        print(_center_text("Reset to default."))
                    else:
                        print(_center_text("Already default."))
                elif val:
                    profile["custom_title"] = val
                    update_full_scan(summary_id, data)
                    print(_center_text("Saved."))
                else:
                    print(_center_text("No change."))

            elif choice == "3":
                # Calculate default summary
                effective_title = profile.get("custom_title")
                if not effective_title:
                    effective_title = "Software Contributor"
                    dev_keywords = {"Development", "Programming", "Engineering"}
                    if any(any(k in s for k in dev_keywords) for s in skills):
                        effective_title = "Software Developer"

                default_summary = f"{effective_title} with a track record of contributions across {len(user_projects)} project(s)."
                if skills:
                    top_skills = skills[:3]
                    default_summary += f" Proficient in {', '.join(top_skills)}"
                    if len(skills) > 3:
                        default_summary += f", along with expertise in {len(skills)-3} other technologies"
                    default_summary += "."

                curr = profile.get("custom_summary", default_summary)
                print(_center_text("Edit Summary (type 'RESET' to restore default):"))
                val = _input_with_prefill("Value: ", curr).strip()

                if val == "RESET":
                    if "custom_summary" in profile:
                        del profile["custom_summary"]
                        update_full_scan(summary_id, data)
                        print(_center_text("Reset to default."))
                    else:
                        print(_center_text("Already default."))
                elif val:
                    profile["custom_summary"] = val
                    update_full_scan(summary_id, data)
                    print(_center_text("Saved."))
                else:
                    print(_center_text("No change."))

            elif choice == "4":
                if not user_projects:
                    print(_center_text("No projects for this user."))
                    continue

                while True:
                    print()
                    print(_center_text(f"Projects for {user}:"))
                    for i, p in enumerate(user_projects, 1):
                        has_custom = " *" if p.get("custom_description") else ""
                        print(_center_text(f"{i}. {p.get('name', 'Unknown')}{has_custom}"))

                    p_sel = input(_center_text("Select project (0 to back): ")).strip()
                    if not p_sel.isdigit():
                        continue
                    
                    p_idx = int(p_sel) - 1
                    if p_idx == -1:
                        break
                    if p_idx < 0 or p_idx >= len(user_projects):
                        continue

                    target_p = user_projects[p_idx]
                    p_name = target_p.get("name")

                    p_context = project_map.get(p_name, {})

                    while True:
                        print()
                        print(_center_text(f"--- Editing Project: {p_name} ---"))
                        
                        sub_opts = [
                            "1. Edit Description",
                            "2. Edit Skills",
                            "0. Back to Project List"
                        ]
                        max_sub = max(len(o) for o in sub_opts)
                        for opt in sub_opts:
                            print(_center_text(opt.ljust(max_sub)))

                        sub_choice = input(_center_text("Choose option: ")).strip()

                        if sub_choice == "0":
                            break

                        if sub_choice == "1":
                            # Reconstruct stats for preview
                            user_stats = {
                                "user_code_files": target_p.get("user_code_files", 0),
                                "user_test_files": target_p.get("user_test_files", 0),
                                "user_doc_files": target_p.get("user_doc_files", 0),
                                "user_design_files": target_p.get("user_design_files", 0),
                                "pct": target_p.get("pct", 0.0),
                                "score": target_p.get("score", 0.0),
                                "files_worked": target_p.get("files_worked", 0),
                                "commit_count": target_p.get("commit_count", 0)
                            }
                            if user_stats["files_worked"] == 0 and target_p.get("files_list"):
                                user_stats["files_worked"] = len(target_p.get("files_list"))

                            default_desc = _build_personal_project_description(p_name, p_context, user_stats)
                            current_custom = target_p.get("custom_description")

                            print("\n" + "="*60)
                            print(f"Project: {p_name}")
                            print(f"Default Generated: {default_desc}")
                            if current_custom:
                                print(f"Current Custom:    {current_custom}")
                            print("="*60)

                            print(_center_text("Edit description (type 'RESET' to restore default):"))
                            prefill = current_custom if current_custom else default_desc
                            new_desc = _input_with_prefill("Value: ", prefill).strip()

                            if new_desc == "RESET":
                                if "custom_description" in target_p:
                                    del target_p["custom_description"]
                                    print(_center_text("Reset to default."))
                                    update_full_scan(summary_id, data)
                            elif new_desc:
                                target_p["custom_description"] = new_desc
                                print(_center_text("Saved custom description."))
                                update_full_scan(summary_id, data)
                            else:
                                print(_center_text("No change."))

                        elif sub_choice == "2":
                            pcs = p_context.get("per_contributor_skills", {})
                            default_skills = pcs.get(user, [])
                            default_skills_str = ", ".join(default_skills)

                            current_custom = target_p.get("custom_skills")
                            current_str = ", ".join(current_custom) if current_custom is not None else default_skills_str

                            print(_center_text("Edit Skills (comma-separated, type 'RESET' to restore default):"))
                            val = _input_with_prefill("Value: ", current_str).strip()

                            if val == "RESET":
                                if "custom_skills" in target_p:
                                    del target_p["custom_skills"]
                                    update_full_scan(summary_id, data)
                                    print(_center_text("Reset to default."))
                                else:
                                    print(_center_text("Already default."))
                            else:
                                # Parse list
                                new_skills = [s.strip() for s in val.split(",") if s.strip()]
                                target_p["custom_skills"] = new_skills
                                update_full_scan(summary_id, data)
                                print(_center_text("Saved."))

            elif choice == "5":
                out = generate_contributor_portfolio(
                    user,
                    profile,
                    project_map,
                    scan_timestamp=scan["timestamp"],
                    sort_mode="date"
                )
                if out:
                    print(_center_text(f"Saved resume to:\n{out}"))
                input(_center_text("Press Enter to continue..."))

            elif choice == "6":
                if get_yes_no(f"Are you sure you want to discard ALL manual edits for {user}?"):
                    # Clear profile-level fields
                    # Remove any key starting with 'custom_' to ensure complete reset
                    for k in list(profile.keys()):
                        if k.startswith("custom_"):
                            profile.pop(k, None)
                    
                    # Clear project-level fields
                    for p in user_projects:
                        for k in list(p.keys()):
                            if k.startswith("custom_") and k != "custom_portfolio_description":
                                p.pop(k, None)
                    
                    update_full_scan(summary_id, data)
                    print(_center_text("All custom fields cleared. Reverted to defaults."))
                    print(_center_text("Resume fields cleared."))